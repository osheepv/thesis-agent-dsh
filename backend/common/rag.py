# -*- coding: utf-8 -*-
"""RAG：本地向量 + 关键词混合检索（零成本方案）。

设计（对齐爸爸决策：本地嵌入模型 + numpy 向量，不用 Chroma/Faiss）：
    - 嵌入：BAAI/bge-small-zh-v1.5（约 100MB，装好永久免费，用户本地 CPU 出算力）。
    - 检索：numpy 余弦 + 轻量BM25；嵌入模型不可用时仍可关键词检索。
    - 数据：用户知识库文件（files/*.pdf 等）提取全文 → 分块（500 字/块）→ 嵌入
      → 持久化 storage/kb/{session_id}/vectors.json（增量：块哈希不变跳过）。
    - 降级：模型不可用 / 无文件 → 返回空结果（写作提示"知识库内容不可用"，不阻塞）。

使用：
    from common.rag import search_kb_blocks
    blocks = search_kb_blocks(session_id, "深度学习加速方法", k=6)
"""
from __future__ import annotations

import hashlib
import json
import logging
import math
import os
import re
import threading
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger("thesis.rag")

#: 分块大小（字）与重叠
_CHUNK_SIZE = 500
_CHUNK_OVERLAP = 50
#: 默认取 Top-K 块
_DEFAULT_K = 6
#: 相似度阈值（低于此值的块视为无关，不注入）
_DEFAULT_MIN_SCORE = 0.25

#: 模型名（缓存于 HF 本地目录）
_MODEL_NAME = "BAAI/bge-small-zh-v1.5"

#: 环境开关（false = 完全禁用 RAG，写作不注入知识库内容）
_ENV_ENABLED = "THESIS_RAG_ENABLED"


def rag_enabled() -> bool:
    return os.getenv(_ENV_ENABLED, "true").lower() != "false"


# ---------------------------------------------------------------------
# 单例模型（懒加载一次，进程内复用）
# ---------------------------------------------------------------------
_model = None
_model_lock = threading.Lock()


def _get_model():
    """加载嵌入模型（单例）。失败返回 None（RAG 降级）。"""
    global _model
    if _model is None:
        with _model_lock:
            if _model is None:
                try:
                    from sentence_transformers import SentenceTransformer

                    logger.info("加载嵌入模型 %s（首次约 4s）", _MODEL_NAME)
                    _model = SentenceTransformer(_MODEL_NAME, device="cpu")
                except Exception as exc:  # noqa: BLE001
                    logger.warning("嵌入模型不可用（RAG 降级）: %s", exc)
                    return None
    return _model


# ---------------------------------------------------------------------
# 文本提取 / 分块
# ---------------------------------------------------------------------
def _extract_text(path: str) -> str:
    """从文件提取纯文本（支持pdf/docx/txt/md）。"""
    ext = Path(path).suffix.lower()
    try:
        if ext == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(path)
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        if ext == ".docx":
            from docx import Document

            document = Document(path)
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
            for table in document.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells if cell.text)
            return "\n".join(parts)
        if ext in (".txt", ".md"):
            return Path(path).read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:  # noqa: BLE001
        logger.warning("提取文本失败 %s: %s", path, exc)
    return ""


def _split_chunks(text: str, size: int = _CHUNK_SIZE, overlap: int = _CHUNK_OVERLAP) -> List[str]:
    """按固定长度切块（带重叠，跨块语义不断裂）。"""
    text = re.sub(r"\s+", " ", text or "").strip()
    if not text:
        return []
    chunks: List[str] = []
    step = max(size - overlap, size // 2)
    for i in range(0, len(text), step):
        chunk = text[i:i + size]
        if chunk:
            chunks.append(chunk)
    return chunks


# ---------------------------------------------------------------------
# 向量库（JSON 落盘 + 增量）
# ---------------------------------------------------------------------
def _vectors_path(session_id: str) -> Path:
    from knowledge.store import _session_dir

    return _session_dir(session_id) / "vectors.json"


def _load_vectors(session_id: str) -> Dict[str, Any]:
    p = _vectors_path(session_id)
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:  # noqa: BLE001
            logger.warning("vectors.json 损坏，重建")
    return {"blocks": [], "hashes": {}}


def _save_vectors(session_id: str, data: Dict[str, Any]) -> None:
    _vectors_path(session_id).write_text(
        json.dumps(data, ensure_ascii=False), encoding="utf-8"
    )


def _hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _keyword_tokens(text: str) -> List[str]:
    """中英文轻量分词：英文单词/编号 + 中文2字gram。"""
    normalized = (text or "").casefold()
    tokens = re.findall(r"[a-z0-9][a-z0-9_./-]{1,}", normalized)
    for sequence in re.findall(r"[一-鿿]{2,}", normalized):
        if len(sequence) <= 12:
            tokens.append(sequence)
        tokens.extend(
            sequence[index:index + 2]
            for index in range(len(sequence) - 1)
        )
    return tokens


def _bm25_scores(query: str, texts: List[str]) -> List[float]:
    """对少量本地块即时计算BM25，不引入额外索引依赖。"""
    query_tokens = list(dict.fromkeys(_keyword_tokens(query)))
    documents = [_keyword_tokens(text) for text in texts]
    if not query_tokens or not documents:
        return [0.0] * len(texts)
    average_length = sum(len(tokens) for tokens in documents) / max(len(documents), 1)
    document_frequency = {
        token: sum(token in set(document) for document in documents)
        for token in query_tokens
    }
    scores: List[float] = []
    k1, b = 1.5, 0.75
    for document in documents:
        counts = Counter(document)
        length_norm = 1 - b + b * len(document) / max(average_length, 1.0)
        score = 0.0
        for token in query_tokens:
            frequency = counts.get(token, 0)
            if not frequency:
                continue
            idf = math.log(
                1 + (len(documents) - document_frequency[token] + 0.5)
                / (document_frequency[token] + 0.5)
            )
            score += idf * frequency * (k1 + 1) / (frequency + k1 * length_norm)
        scores.append(score)
    return scores


def _embed(model, texts: List[str]) -> List[List[float]]:
    """批量嵌入（无 stride 开销；CPU 一次几百字很快）。"""
    if not texts:
        return []
    vecs = model.encode(texts, batch_size=32, show_progress_bar=False)
    return [v.tolist() for v in vecs]


# ---------------------------------------------------------------------
# 对外：索引知识库文件
# ---------------------------------------------------------------------
def index_session(session_id: str) -> int:
    """索引会话知识库全部文本文件（增量：哈希不变跳过）。返回新增块数。"""
    if not rag_enabled():
        return 0
    model = _get_model()
    from knowledge.store import get_kb_store

    store = get_kb_store()
    sdir = Path(store.session_path(session_id))
    files_dir = sdir / "files"
    data = _load_vectors(session_id)
    hashes = data.get("hashes", {})
    blocks = data.get("blocks", [])
    added = 0

    if files_dir.exists():
        for f in sorted(files_dir.iterdir()):
            if not f.is_file() or f.suffix.lower() not in (".pdf", ".docx", ".txt", ".md"):
                continue
            try:
                mtime_s = str(int(f.stat().st_mtime))
            except OSError:
                continue
            key = f.name
            existing = [block for block in blocks if block.get("file") == key]
            if (
                hashes.get(key) == mtime_s
                and existing
                and (model is None or all(block.get("vector") for block in existing))
            ):
                continue  # 未变化，跳过
            text = _extract_text(str(f))
            chunks = _split_chunks(text)
            if not chunks:
                hashes[key] = mtime_s
                continue
            vecs = _embed(model, chunks) if model is not None else [[] for _ in chunks]
            # 移除该文件旧块，追加新块
            blocks = [b for b in blocks if b.get("file") != key]
            for cd, cv in zip(chunks, vecs):
                blocks.append({
                    "file": key,
                    "source": str(f),
                    "text": cd,
                    "vector": cv,
                    "hash": _hash_text(cd),
                })
            hashes[key] = mtime_s
            added += len(chunks)
            logger.info("索引 %s：%d 块", key, len(chunks))
    data = {"blocks": blocks, "hashes": hashes}
    _save_vectors(session_id, data)
    return added


# ---------------------------------------------------------------------
# 对外：语义检索
# ---------------------------------------------------------------------
def search_kb_blocks(session_id: str, query: str, k: int = _DEFAULT_K,
                     min_score: float = _DEFAULT_MIN_SCORE) -> List[Dict[str, Any]]:
    """在会话知识库全文块中做混合检索，返回Top-K块（降序）。

    Returns:
        [{file, text, score, vector_score, keyword_score, retrieval_mode}]。
    """
    if not rag_enabled() or not query:
        return []
    data = _load_vectors(session_id)
    blocks = data.get("blocks", [])
    if not blocks and _vectors_path(session_id).exists():
        # 已有向量文件但块为空 → 文件级索引已完成、无内容，无需重索引
        return []
    if not blocks:
        # 懒索引：首次检索时自动索引知识库文件（无向量文件即未索引）
        try:
            index_session(session_id)
            data = _load_vectors(session_id)
            blocks = data.get("blocks", [])
        except Exception as exc:  # noqa: BLE001
            logger.warning("懒索引失败: %s", exc)
    if not blocks:
        return []
    texts = [str(block.get("text", "")) for block in blocks]
    keyword_raw = _bm25_scores(query, texts)
    keyword_max = max(keyword_raw, default=0.0)
    keyword_scores = [
        score / keyword_max if keyword_max > 0 else 0.0
        for score in keyword_raw
    ]
    vector_scores: Optional[List[float]] = None
    model = _get_model()
    try:
        if model is not None and all(block.get("vector") for block in blocks):
            import numpy as np

            qvec = _embed(model, [query])[0]
            matrix = np.array([block["vector"] for block in blocks], dtype=np.float32)
            q = np.array(qvec, dtype=np.float32)
            cosine = (matrix @ q) / (
                np.linalg.norm(matrix, axis=1) * (np.linalg.norm(q) + 1e-9) + 1e-9
            )
            vector_scores = [max(0.0, float(score)) for score in cosine]
    except Exception as exc:  # noqa: BLE001
        logger.warning("向量检索失败，降级关键词检索: %s", exc)
        vector_scores = None

    if vector_scores is not None and keyword_max > 0:
        combined = [
            vector * 0.7 + keyword * 0.3
            for vector, keyword in zip(vector_scores, keyword_scores)
        ]
        mode = "hybrid"
    elif vector_scores is not None:
        combined = vector_scores
        mode = "vector"
    else:
        combined = keyword_scores
        mode = "keyword"
    order = sorted(range(len(blocks)), key=lambda index: combined[index], reverse=True)[:k]
    results = []
    for index in order:
        score = combined[index]
        if score < min_score:
            continue
        block = blocks[index]
        results.append({
            "file": block.get("file", ""),
            "text": block.get("text", ""),
            "score": round(score, 4),
            "vector_score": round(vector_scores[index], 4) if vector_scores is not None else None,
            "keyword_score": round(keyword_scores[index], 4),
            "retrieval_mode": mode,
        })
    return results


def kb_blocks_text(session_id: str, query: str, k: int = _DEFAULT_K,
                   max_chars: int = 3000) -> str:
    """检索并把 Top-K 块拼成注入 prompt 的文本块（环6 写作用）。"""
    hits = search_kb_blocks(session_id, query, k=k)
    if not hits:
        return ""
    lines = ["【知识库相关段落（混合检索，仅可参考，不可编造未给内容）】"]
    budget = max_chars
    for i, h in enumerate(hits, start=1):
        if budget <= 0:
            break
        seg = h["text"][:budget]
        lines.append(f"[KB{i}|{h['file']}|score={h['score']}] {seg}")
        budget -= len(seg)
    return "\n".join(lines)
