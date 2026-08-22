# -*- coding: utf-8 -*-
"""M6 docx 生成器：基于 docxtpl 按用户模板渲染。

设计要点：
    - 读取已解析的模板文件路径 + 占位符列表；
    - 用 Jinja2 变量代入内容映射（topic/outline/chapter 等）；
    - 生成合规 .docx 到输出目录，返回文件路径；
    - 生成后交由 validator 校验，校验不通过则拒绝交付（由上层回退）。

依赖：docxtpl（0.17.x 兼容实现）。
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Any, Dict, List, Mapping, Optional

from docxtpl import DocxTemplate

from common.aicoding.exception import BizException, ErrorCode

from ..config import DocxConfig


@dataclass
class GenerateOutcome:
    """生成结果载体。

    Attributes:
        file_path: 生成文件落盘路径。
        filename: 生成文件名。
        word_count: 估算字数。
        supplied_keys: 内容映射中实际代入的键。
        missing_keys: 模板占位符但内容映射未提供的键（渲染后仍残留 {{}}）。
    """

    file_path: str = ""
    filename: str = ""
    word_count: int = 0
    supplied_keys: List[str] = field(default_factory=list)
    missing_keys: List[str] = field(default_factory=list)


class DocxGenerator:
    """M6 docx 生成器。

    Args:
        config: docx 模块配置。
    """

    def __init__(self, config: Optional[DocxConfig] = None) -> None:
        self._config = config or DocxConfig()

    def render(
        self,
        template_path: str,
        content: Mapping[str, Any],
        template_placeholders: Optional[List[str]] = None,
        output_path: Optional[str] = None,
        filename: Optional[str] = None,
    ) -> GenerateOutcome:
        """按模板渲染生成 docx。

        Args:
            template_path: 模板文件落盘路径。
            content: 占位符 -> 内容映射。
            template_placeholders: 模板占位符列表（用于缺失填充诊断）。
            output_path: 输出文件路径（可选，默认写入 config.OUTPUT_DIR）。
            filename: 生成文件名（可选，默认以 template 名派生）。

        Returns:
            GenerateOutcome。

        Raises:
            BizException: 模板不存在或渲染失败时抛出。
        """
        self._config.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        out_path = output_path or self._default_output_path(filename)

        try:
            tpl = DocxTemplate(template_path)
        except Exception as exc:  # noqa: BLE001
            raise BizException(
                ErrorCode.DOCX_GENERATE_FAILED,
                "模板打开失败",
                detail={"template_path": template_path, "err": str(exc)},
            ) from exc

        context: Dict[str, Any] = dict(content or {})
        try:
            tpl.render(context)
        except Exception as exc:  # noqa: BLE001
            raise BizException(
                ErrorCode.DOCX_GENERATE_FAILED,
                "模板渲染失败（占位符/变量不匹配）",
                detail={"err": str(exc), "supplied": list(context.keys())},
            ) from exc

        try:
            tpl.save(out_path)
        except Exception as exc:  # noqa: BLE001
            raise BizException(
                ErrorCode.DOCX_GENERATE_FAILED,
                "生成文件写入失败",
                detail={"out_path": out_path, "err": str(exc)},
            ) from exc

        word_count = self._estimate_word_count(out_path)
        supplied_keys = list(context.keys())

        missing_keys: List[str] = []
        if template_placeholders:
            missing_keys = [k for k in template_placeholders if k not in context]
        # 二次诊断：渲染后仍残留在 XML 中的 {{ 占位符
        if not missing_keys:
            missing_keys = self._detect_leftover_placeholders(out_path)

        return GenerateOutcome(
            file_path=out_path,
            filename=out_path.rsplit("/", 1)[-1].rsplit("\\", 1)[-1],
            word_count=word_count,
            supplied_keys=supplied_keys,
            missing_keys=missing_keys,
        )

    # ------------------------------------------------------------------ #
    def _default_output_path(self, filename: Optional[str]) -> str:
        """派生默认输出路径。"""
        import uuid

        name = (filename or "thesis_render.docx")
        if not name.lower().endswith(".docx"):
            name += ".docx"
        safe_name = name.replace("/", "_").replace("\\", "_").replace("..", "_")
        return str(self._config.OUTPUT_DIR / f"{uuid.uuid4().hex}_{safe_name}")

    @staticmethod
    def _estimate_word_count(path: str) -> int:
        """估算 DOCX 正文总字数（剔除空白字符）。"""
        try:
            from docx import Document

            doc = Document(path)
            chars = sum(len(p.text.replace(" ", "")) for p in doc.paragraphs)
            return chars
        except Exception:  # noqa: BLE001
            return 0

    @staticmethod
    def _detect_leftover_placeholders(path: str) -> List[str]:
        """检测渲染后仍未替换的 Jinja2 占位符（诊断为缺失变量）。"""
        import re

        pattern = re.compile(r"\{\{\s*([A-Za-z_][\w\.]*)\s*\}\}")
        try:
            from docx import Document

            doc = Document(path)
            text_parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            found: List[str] = []
            for part in text_parts:
                for match in pattern.finditer(part):
                    if match.group(1) not in found:
                        found.append(match.group(1))
            return found
        except Exception:  # noqa: BLE001
            return []
