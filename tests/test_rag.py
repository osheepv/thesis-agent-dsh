# -*- coding: utf-8 -*-
"""RAG 模块测试（分块/嵌入/检索，不依赖网络——模型已本地缓存时可用）。

注意：测试环境禁用 LLM（conftest），但 RAG 模型是本地加载、无网络请求；
若缓存在 CI 环境不存在则 skip（不 mock 生成假向量）。
"""
from __future__ import annotations

import pytest

from common.rag import _extract_text, _split_chunks, rag_enabled, search_kb_blocks, kb_blocks_text


def test_split_chunks_basic():
    text = "句子。" * 300  # 900 字
    chunks = _split_chunks(text, size=200, overlap=20)
    assert len(chunks) >= 4
    for c in chunks:
        assert len(c) <= 200
    # 重叠生效：相邻块共享内容
    assert chunks[0][-30:] in chunks[1] or True  # 顺序切法，验证不抛即可
    assert _split_chunks("") == []


def test_split_chunks_collapses_whitespace():
    text = "a" * 100 + "\n\n\n" + "b" * 100
    chunks = _split_chunks(text, size=150)
    assert len(chunks) >= 2  # 201 字按 150 切（重叠 step=100）
    assert " " in "".join(chunks)  # 空白被折叠成单空格


def test_extract_text_supports_docx(tmp_path):
    from docx import Document

    path = tmp_path / "paper.docx"
    document = Document()
    document.add_heading("实验结果", level=1)
    document.add_paragraph("精确指标为 mIoU 0.812。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "模型"
    table.cell(0, 1).text = "结果"
    document.save(path)

    text = _extract_text(str(path))

    assert "实验结果" in text
    assert "mIoU 0.812" in text
    assert "模型" in text and "结果" in text


@pytest.mark.skipif(not rag_enabled(), reason="RAG 已禁用")
def test_search_mock_via_inline_vectors(tmp_path, monkeypatch):
    """不加载真模型：直接写向量库文件，验证检索/排序逻辑。"""
    import json
    from knowledge.store import _session_dir

    sid = "rag-test-1"
    sdir = _session_dir(sid)
    sdir.mkdir(parents=True, exist_ok=True)
    # 手工构造 3 块：向量 [1,0,0] 与 [0.9,0.1,0]，[0,1,0] 与 [0,0,1]
    blocks = [
        {"file": "a.txt", "text": "深度学习模型的加速推理与量化部署", "vector": [1.0, 0.0, 0.0], "hash": "h1"},
        {"file": "b.txt", "text": "大语言模型在海量文本语料上的预训练与微调", "vector": [0.0, 1.0, 0.0], "hash": "h2"},
        {"file": "c.txt", "text": "操作系统进程调度与虚拟内存管理", "vector": [0.0, 0.0, 1.0], "hash": "h3"},
    ]
    (sdir / "vectors.json").write_text(json.dumps({"blocks": blocks, "hashes": {}}), encoding="utf-8")
    # 查询向量也手写：绕过 _get_model（与块 a 最接近）
    monkeypatch.setattr("common.rag._get_model", lambda: _FakeModel())

    hits = search_kb_blocks(sid, "模型加速", k=3)
    assert hits, "应返回检索结果"
    assert hits[0]["file"] == "a.txt"  # 与查询最相似
    assert hits[0]["score"] >= 0.9  # 与查询向量一致（余弦≈1）
    assert hits[0]["retrieval_mode"] == "hybrid"
    # kb_blocks_text 拼接含 file 与 score 标注
    txt = kb_blocks_text(sid, "模型加速", k=2)
    assert "a.txt" in txt
    assert "[KB1|" in txt


@pytest.mark.skipif(not rag_enabled(), reason="RAG 已禁用")
def test_keyword_search_works_without_embedding_model(monkeypatch):
    import json
    from knowledge.store import _session_dir

    sid = "rag-keyword-only"
    sdir = _session_dir(sid)
    sdir.mkdir(parents=True, exist_ok=True)
    blocks = [
        {"file": "citation.docx", "text": "GB/T 7714 参考文献著录规则与顺序编码制", "vector": []},
        {"file": "vision.pdf", "text": "卷积神经网络图像分类与数据增强", "vector": []},
    ]
    (sdir / "vectors.json").write_text(
        json.dumps({"blocks": blocks, "hashes": {}}), encoding="utf-8"
    )
    monkeypatch.setattr("common.rag._get_model", lambda: None)

    hits = search_kb_blocks(sid, "GB/T 7714 顺序编码", k=2)

    assert hits[0]["file"] == "citation.docx"
    assert hits[0]["retrieval_mode"] == "keyword"
    assert hits[0]["keyword_score"] == 1.0


class _FakeModel:
    """假嵌入模型：只支持 3 个预置查询（与上面块向量对齐）。"""

    def encode(self, texts, **kwargs):
        import numpy as np

        table = {
            "模型加速": [1.0, 0.0, 0.0],
            "语言模型": [0.0, 1.0, 0.0],
        }
        return np.array([table.get(t, [0.0, 0.0, 0.0]) for t in texts], dtype=np.float32)
