# -*- coding: utf-8 -*-
"""M6 docx 校验器：基于 openxml-audit 做 OOXML 校验。

校验维度（对应 evidence ladder 的“校验是地板层”）：
    1) schema/semantic 校验：ECMA 合法性（openxml-audit OpenXmlValidator）。
    2) load 校验：能否被 python-docx 正常重开（近似“目标应用可加载”）。
    3) roundtrip 校验：加载后原样再保存一次、重载不报错（近似“重存不被重写意图”）。

本模块只保留校验结果，不深究内部修复。校验不通过时上层拒绝交付/触发回退。
"""
from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

from common.aicoding.exception import BizException, ErrorCode

from ..config import DocxConfig
from ..cross_reference import audit_cross_references


@dataclass
class ValidateOutcome:
    """校验结果内部载体。

    Attributes:
        is_valid: 整体是否通过（schema+load+roundtrip 全通过）。
        schema_valid: schema/semantic 校验是否通过。
        load_valid: load（能否重开）是否通过。
        roundtrip_valid: roundtrip（保存重载）是否通过。
        error_count / warning_count: 统计。
        errors / warnings: openxml_audit 抽取明细。
        validator: 校验器版本。
    """

    is_valid: bool = False
    schema_valid: bool = False
    load_valid: bool = False
    roundtrip_valid: bool = False
    error_count: int = 0
    warning_count: int = 0
    errors: List[Dict[str, Any]] = field(default_factory=list)
    warnings: List[Dict[str, Any]] = field(default_factory=list)
    validator: str = "openxml-audit"
    cross_reference_valid: bool = True
    cross_reference_report: Dict[str, Any] = field(default_factory=dict)


class DocxValidator:
    """M6 docx 校验器。

    Args:
        config: docx 模块配置。
    """

    def __init__(self, config: Optional[DocxConfig] = None) -> None:
        self._config = config or DocxConfig()
        self._version = self._detect_version()

    # ------------------------------------------------------------------ #
    # 对外入口
    # ------------------------------------------------------------------ #
    def validate(self, file_path: str, strict: bool = True) -> ValidateOutcome:
        """对 docx 执行多层校验并返回结果。

        Args:
            file_path: 待校验文件路径。
            strict: 是否启用严格 schema/semantic 校验。

        Returns:
            ValidateOutcome。

        Raises:
            BizException: 文件不存在或校验器初始化失败时抛出。
        """
        import os

        if not os.path.exists(file_path):
            raise BizException(
                ErrorCode.DOCX_VALIDATE_FAILED, "待校验文件不存在", detail={"path": file_path}
            )

        # 维度1：schema/semantic
        schema = self._schema_validate(file_path, strict=strict)
        errors: List[Dict[str, Any]] = list(schema["errors"])
        warnings: List[Dict[str, Any]] = list(schema["warnings"])
        schema_valid = schema["is_valid"]
        error_count = schema["error_count"]
        warning_count = schema["warning_count"]

        # 维度2：load
        load_valid = self._load_validate(file_path)

        # 维度3：roundtrip
        roundtrip_valid = self._roundtrip_validate(file_path)

        cross_reference_report = audit_cross_references(file_path).to_dict()
        cross_reference_valid = bool(cross_reference_report["passed"])
        if not cross_reference_valid:
            errors.append(
                {
                    "severity": "ERROR",
                    "description": "DOCX 书签/REF 交叉引用不完整",
                    "part": "/word/document.xml",
                    "detail": cross_reference_report,
                }
            )
            error_count += 1

        is_valid = schema_valid and load_valid and roundtrip_valid and cross_reference_valid
        return ValidateOutcome(
            is_valid=is_valid,
            schema_valid=schema_valid,
            load_valid=load_valid,
            roundtrip_valid=roundtrip_valid,
            error_count=error_count,
            warning_count=warning_count,
            errors=errors,
            warnings=warnings,
            validator=self._version,
            cross_reference_valid=cross_reference_valid,
            cross_reference_report=cross_reference_report,
        )

    # ------------------------------------------------------------------ #
    # 维度实现
    # ------------------------------------------------------------------ #
    def _schema_validate(self, file_path: str, strict: bool = True) -> Dict[str, Any]:
        """用 openxml_audit 做 schema/semantic 校验。"""
        import openxml_audit as oa

        try:
            validator = oa.OpenXmlValidator(
                file_format=oa.FileFormat.OFFICE_2019,
                max_errors=self._config.VALIDATE_MAX_ERRORS,
                schema_validation=True,
                semantic_validation=strict,
            )
            result = validator.validate(file_path)
        except Exception as exc:  # noqa: BLE001
            # 校验器本身报错视为“未能证明合法”，按不通过处理
            return {
                "is_valid": False,
                "error_count": 1,
                "warning_count": 0,
                "errors": [{"severity": "ERROR", "description": f"校验器异常: {exc}", "part": ""}],
                "warnings": [],
            }

        errors = [self._serialize_error(e) for e in result.errors[: self._config.VALIDATE_MAX_ERRORS]]
        warnings = [
            self._serialize_error(e) for e in getattr(result, "warnings", [])[: self._config.VALIDATE_MAX_ERRORS]
        ]
        return {
            "is_valid": bool(result.is_valid),
            "error_count": int(result.error_count),
            "warning_count": int(result.warning_count),
            "errors": errors,
            "warnings": warnings,
        }

    @staticmethod
    def _load_validate(file_path: str) -> bool:
        """load 校验：python-docx 能否正常重开（弱化代理“目标应用可打开”）。"""
        try:
            from docx import Document

            Document(file_path)
            return True
        except Exception:  # noqa: BLE001
            return False

    @staticmethod
    def _roundtrip_validate(file_path: str) -> bool:
        """roundtrip 校验：加载后原样保存到临时文件并重载，不报错即通过。"""
        import tempfile

        try:
            from docx import Document

            doc = Document(file_path)
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            try:
                doc.save(tmp_path)
                Document(tmp_path)  # 重载
                return True
            finally:
                import os

                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        except Exception:  # noqa: BLE001
            return False

    # ------------------------------------------------------------------ #
    @staticmethod
    def _serialize_error(err: Any) -> Dict[str, Any]:
        """将 openxml_audit 的 ValidationError 序列化为字典。"""
        severity = getattr(err, "severity", None)
        return {
            "severity": severity.value if hasattr(severity, "value") else str(severity),
            "description": getattr(err, "description", "") or "",
            "part": getattr(err, "part_uri", "") or "",
            "id": getattr(err, "id", "") or "",
            "node": str(getattr(err, "node", "") or "")[:300],
        }

    @staticmethod
    def _detect_version() -> str:
        """探测 openxml_audit 版本。"""
        try:
            import openxml_audit as oa

            return f"openxml-audit-{getattr(oa, '__version__', '?')}"
        except Exception:  # noqa: BLE001
            return "openxml-audit"
