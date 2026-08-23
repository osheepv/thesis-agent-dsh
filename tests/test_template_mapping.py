"""学校模板持久化、占位符映射与真实 DOCX 渲染测试。"""

from __future__ import annotations

import io
import json
from pathlib import Path

from docx import Document

from application.service.uc_main_orchestration import (
    MainOrchestration,
    RealDocxRenderer,
    TaskRecord,
)
from common.aicoding.enums import Degree
from executor.base import ExecResult
from thesis_docx.config import DocxConfig
from thesis_docx.service import DocxService


def _template_bytes() -> bytes:
    document = Document()
    document.add_paragraph("学校题目：{{ 学校题目 }}")
    document.add_paragraph("学校正文：{{ 学校正文 }}")
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


class _Executor:
    def __init__(self, ring_no: int) -> None:
        self.ring_no = ring_no

    def execute(self, ctx) -> ExecResult:
        payloads = {
            1: {"candidates": [{"title": "学校模板映射研究"}], "recommendation": "推荐"},
            2: {"novelty_level": "HIGH", "similar_count": 0, "recommendation": "通过"},
            3: {"items": [{"title": "真实文献", "doi": "10.1000/template"}], "summary": "1条"},
            4: {"verdict": "顺", "overlap_count": 0, "recommendation": "通过"},
            5: {"theme": "学校模板映射研究", "chapters": [{"level": 1, "number": "1", "title": "绪论"}]},
            6: {"chapters": [{"chapter_no": 1, "chapter_title": "绪论", "content": "正文 [L1]", "word_count": 7}], "total_words": 7, "used_refs": ["[L1]"]},
            7: {"chapters": [{"chapter_no": 1, "chapter_title": "绪论", "content": "润色正文 [L1]", "word_count": 9}], "total_words": 9},
            8: {"total": 1, "passed": 1, "uncertain": 0, "failed": 0, "summary": "通过"},
        }
        return ExecResult(
            output=json.dumps(payloads[self.ring_no], ensure_ascii=False),
            accept=True,
            evidence={"source": "test-double"},
        )


def test_task_record_persists_template_configuration():
    record = TaskRecord(
        "task", "title", "MASTER", "field",
        template_id="TPL-1",
        template_path="C:/templates/school.docx",
        template_name="school.docx",
        template_placeholders=["学校题目", "学校正文"],
        template_mapping={"学校题目": "title", "学校正文": "content"},
    )
    restored = TaskRecord.from_dict(record.to_dict())
    assert restored.template_path.endswith("school.docx")
    assert restored.template_name == "school.docx"
    assert restored.template_placeholders == ["学校题目", "学校正文"]
    assert restored.template_mapping["学校正文"] == "content"


def test_uploaded_school_template_is_used_for_final_render(tmp_path, monkeypatch):
    monkeypatch.setattr(
        "application.service.uc_main_orchestration.get_executor",
        lambda ring_no: _Executor(int(ring_no)),
    )
    config = DocxConfig()
    config.UPLOAD_DIR = tmp_path / "templates"
    config.OUTPUT_DIR = tmp_path / "outputs"
    service = DocxService(config=config)
    renderer = RealDocxRenderer(repository=service._repo)  # noqa: SLF001
    renderer._parser = service._parser  # noqa: SLF001
    renderer._generator = service._generator  # noqa: SLF001
    orchestration = MainOrchestration(docx_renderer=renderer)
    orchestration._docx_service = service  # noqa: SLF001

    task_id = orchestration.create_task(
        "初始题目", Degree.MASTER, "人工智能", session_id="template-test"
    ).data["task_id"]
    uploaded = orchestration.upload_template(
        task_id, _template_bytes(), "学校模板.docx"
    ).data
    assert uploaded["mapping"] == {
        "学校题目": "title",
        "学校正文": "content",
    }
    config_result = orchestration.get_template_config(task_id).data
    assert config_result["is_custom"] is True
    assert config_result["placeholders"] == ["学校题目", "学校正文"]

    for ring_no, runner in (
        (1, orchestration.run_ring1), (2, orchestration.run_ring2),
        (3, orchestration.run_ring3), (4, orchestration.run_ring4),
        (5, orchestration.run_ring5), (6, orchestration.run_ring6),
        (7, orchestration.run_ring7), (8, orchestration.run_ring8),
    ):
        assert runner(task_id).is_ok
        orchestration.confirm_ring(task_id, ring_no)

    generated = orchestration.generate_docx(task_id).data
    output = Document(generated["file_path"])
    text = "\n".join(paragraph.text for paragraph in output.paragraphs)
    assert "学校题目：学校模板映射研究" in text
    assert "学校正文：" in text
    assert "润色正文 [L1]" in text
    assert "{{ 学校题目 }}" not in text


def test_template_mapping_rejects_unknown_target(monkeypatch):
    orchestration = MainOrchestration()
    task_id = orchestration.create_task(
        "映射校验", Degree.MASTER, "人工智能", session_id="template-invalid"
    ).data["task_id"]
    record = orchestration._store.get(task_id)  # noqa: SLF001
    record.template_id = "TPL-1"
    record.template_path = "C:/templates/school.docx"
    record.template_placeholders = ["学校正文"]
    orchestration._store.put(record)  # noqa: SLF001
    try:
        orchestration.set_template_mapping(task_id, {"不存在": "content"})
    except Exception as exc:  # noqa: BLE001
        assert "不存在的占位符" in str(exc)
    else:
        raise AssertionError("未知占位符必须被拒绝")


def test_template_mapping_console_api(tmp_path):
    from application.main import build_app
    from fastapi.testclient import TestClient

    config = DocxConfig()
    config.UPLOAD_DIR = tmp_path / "templates"
    config.OUTPUT_DIR = tmp_path / "outputs"
    service = DocxService(config=config)
    orchestration = MainOrchestration()
    orchestration._docx_service = service  # noqa: SLF001
    task_id = orchestration.create_task(
        "模板 API", Degree.MASTER, "人工智能", session_id="template-api"
    ).data["task_id"]
    client = TestClient(build_app(orchestration=orchestration))
    uploaded = client.post(
        f"/api/v1/console/tasks/{task_id}/template?session_id=template-api",
        files={
            "file": (
                "school.docx",
                _template_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    ).json()
    assert uploaded["code"] == 0
    assert uploaded["data"]["placeholders"] == ["学校题目", "学校正文"]
    saved = client.post(
        f"/api/v1/console/tasks/{task_id}/template/mapping?session_id=template-api",
        json={"mapping": {"学校题目": "title", "学校正文": "content"}},
    ).json()
    assert saved["code"] == 0
    detail = client.get(
        f"/api/v1/console/tasks/{task_id}/template?session_id=template-api"
    ).json()
    assert detail["data"]["mapping"]["学校正文"] == "content"
    template_path = Path(orchestration._store.get(task_id).template_path)  # noqa: SLF001
    assert template_path.is_file()
    deleted = orchestration.delete_task(task_id)
    assert deleted.data["template_deleted"] is True
    assert not template_path.exists()
