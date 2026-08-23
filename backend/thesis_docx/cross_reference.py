"""DOCX 原生书签与 REF 域交叉引用。

输入正文使用两种内部标记：

``[[BOOKMARK:TABLE-4-1|表4-1 实验结果]]`` 定义目标；
``[[REF:TABLE-4-1|表4-1]]`` 插入 Word 原生 REF 域。

该层只处理明确标记，不猜测普通文本中的“见表X”，避免把错误文本静默包装成
有效交叉引用。
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


_MARKER_RE = re.compile(
    r"\[\[(?P<kind>BOOKMARK|REF):(?P<target>[A-Za-z][A-Za-z0-9_.-]{0,127})"
    r"\|(?P<label>[^\[\]\r\n]+)\]\]"
)
_ANY_MARKER_RE = re.compile(r"\[\[(?:BOOKMARK|REF):")
_REF_INSTRUCTION_RE = re.compile(r"\bREF\s+([A-Za-z_][A-Za-z0-9_]{0,39})\b")


class CrossReferenceError(ValueError):
    """书签/REF 标记无效或生成后的域不完整。"""


@dataclass(frozen=True)
class CrossReferenceReport:
    bookmark_count: int = 0
    reference_count: int = 0
    bookmark_map: dict[str, str] = field(default_factory=dict)
    unresolved_targets: tuple[str, ...] = ()
    leftover_markers: tuple[str, ...] = ()

    @property
    def passed(self) -> bool:
        return not self.unresolved_targets and not self.leftover_markers

    def to_dict(self) -> dict[str, object]:
        return {
            "passed": self.passed,
            "bookmark_count": self.bookmark_count,
            "reference_count": self.reference_count,
            "bookmark_map": dict(self.bookmark_map),
            "unresolved_targets": list(self.unresolved_targets),
            "leftover_markers": list(self.leftover_markers),
        }


def normalize_bookmark_name(target: str) -> str:
    """把业务目标 ID 转为 Word 允许的、稳定且不超过 40 字符的名称。"""
    normalized = normalize_target_id(target).replace("-", "_").replace(".", "_")
    normalized = f"XREF_{normalized}"
    if len(normalized) <= 40:
        return normalized
    digest = hashlib.sha256(target.encode("utf-8")).hexdigest()[:10]
    return f"{normalized[:29]}_{digest}"


def normalize_target_id(value: str) -> str:
    """把外部表/图/结果 ID 规范为内部标记允许的稳定目标 ID。"""
    raw = value.strip()
    normalized = re.sub(r"[^A-Za-z0-9_.-]", "_", raw)
    if not normalized or not normalized[0].isalpha():
        normalized = f"R_{normalized}"
    if len(normalized) <= 128:
        return normalized
    digest = hashlib.sha256(raw.encode("utf-8")).hexdigest()[:12]
    return f"{normalized[:115]}_{digest}"


def apply_cross_references(path: str | Path) -> CrossReferenceReport:
    """把文档内的内部标记转换为原生 OOXML 书签和 REF 域并验证。"""
    docx_path = str(path)
    document = Document(docx_path)
    paragraphs = list(_iter_document_paragraphs(document))
    definitions: dict[str, str] = {}
    references: list[str] = []
    for paragraph in paragraphs:
        for match in _MARKER_RE.finditer(paragraph.text):
            target = match.group("target")
            if match.group("kind") == "BOOKMARK":
                if target in definitions:
                    raise CrossReferenceError(f"交叉引用目标重复定义: {target}")
                definitions[target] = match.group("label").strip()
            else:
                references.append(target)
        if _ANY_MARKER_RE.search(paragraph.text) and not _MARKER_RE.search(paragraph.text):
            raise CrossReferenceError(f"交叉引用标记格式非法: {paragraph.text}")

    missing = sorted(set(references) - set(definitions))
    if missing:
        raise CrossReferenceError(f"REF 引用了未定义目标: {missing}")
    bookmark_map = {target: normalize_bookmark_name(target) for target in definitions}
    if len(set(bookmark_map.values())) != len(bookmark_map):
        raise CrossReferenceError("书签规范化后发生名称冲突")

    next_bookmark_id = _next_bookmark_id(document)
    for paragraph in paragraphs:
        text = paragraph.text
        if not _MARKER_RE.search(text):
            continue
        tokens: list[tuple[str, str, str]] = []
        cursor = 0
        for match in _MARKER_RE.finditer(text):
            if match.start() > cursor:
                tokens.append(("TEXT", "", text[cursor:match.start()]))
            tokens.append(
                (match.group("kind"), match.group("target"), match.group("label").strip())
            )
            cursor = match.end()
        if cursor < len(text):
            tokens.append(("TEXT", "", text[cursor:]))
        _clear_paragraph_runs(paragraph)
        for kind, target, label in tokens:
            if kind == "TEXT":
                if label:
                    paragraph.add_run(label)
            elif kind == "BOOKMARK":
                _append_bookmark(
                    paragraph,
                    bookmark_id=next_bookmark_id,
                    bookmark_name=bookmark_map[target],
                    label=label,
                )
                next_bookmark_id += 1
            else:
                _append_ref_field(
                    paragraph,
                    bookmark_name=bookmark_map[target],
                    fallback_label=label,
                )

    _enable_update_fields(document)
    document.save(docx_path)
    report = audit_cross_references(docx_path, expected_targets=bookmark_map)
    if not report.passed:
        raise CrossReferenceError(
            "DOCX 交叉引用验证失败: "
            f"unresolved={list(report.unresolved_targets)}, "
            f"leftover={list(report.leftover_markers)}"
        )
    return report


def audit_cross_references(
    path: str | Path,
    *,
    expected_targets: dict[str, str] | None = None,
) -> CrossReferenceReport:
    """读取 DOCX 的 OOXML，核验每个 REF 都有目标书签且无内部标记残留。"""
    document = Document(str(path))
    bookmark_names: set[str] = set()
    ref_names: list[str] = []
    leftovers: list[str] = []
    for paragraph in _iter_document_paragraphs(document):
        if _ANY_MARKER_RE.search(paragraph.text):
            leftovers.append(paragraph.text)
        for node in paragraph._p.iter():
            if node.tag == qn("w:bookmarkStart"):
                name = node.get(qn("w:name"), "")
                if name:
                    bookmark_names.add(name)
            elif node.tag == qn("w:instrText"):
                match = _REF_INSTRUCTION_RE.search(node.text or "")
                if match:
                    ref_names.append(match.group(1))
    unresolved = tuple(sorted(set(ref_names) - bookmark_names))
    reverse = {name: target for target, name in (expected_targets or {}).items()}
    reported_map = {
        reverse.get(name, name): name
        for name in sorted(bookmark_names)
    }
    return CrossReferenceReport(
        bookmark_count=len(bookmark_names),
        reference_count=len(ref_names),
        bookmark_map=reported_map,
        unresolved_targets=unresolved,
        leftover_markers=tuple(leftovers),
    )


def _iter_document_paragraphs(document: DocumentType) -> Iterator[Paragraph]:
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from container.paragraphs
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    visited_cells: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            cell_key = id(cell._tc)
            if cell_key in visited_cells:
                continue
            visited_cells.add(cell_key)
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _next_bookmark_id(document: DocumentType) -> int:
    values: list[int] = []
    for node in document.element.body.iter(qn("w:bookmarkStart")):
        raw = node.get(qn("w:id"), "")
        if raw.isdigit():
            values.append(int(raw))
    return max(values, default=-1) + 1


def _clear_paragraph_runs(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _append_bookmark(
    paragraph: Paragraph,
    *,
    bookmark_id: int,
    bookmark_name: str,
    label: str,
) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    paragraph._p.append(start)
    paragraph.add_run(label)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def _append_ref_field(
    paragraph: Paragraph,
    *,
    bookmark_name: str,
    fallback_label: str,
) -> None:
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run.append(begin)
    paragraph._p.append(begin_run)

    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" REF {bookmark_name} \\h "
    instruction_run.append(instruction)
    paragraph._p.append(instruction_run)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    paragraph._p.append(separate_run)

    display_run = OxmlElement("w:r")
    display_text = OxmlElement("w:t")
    display_text.text = fallback_label
    display_run.append(display_text)
    paragraph._p.append(display_run)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.append(end_run)


def _enable_update_fields(document: DocumentType) -> None:
    settings = document.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        successors = {
            qn(name)
            for name in (
                "w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr", "w:compat",
                "w:docVars", "w:rsids", "m:mathPr", "w:uiCompat97To2003",
                "w:attachedSchema", "w:themeFontLang", "w:clrSchemeMapping",
                "w:doNotIncludeSubdocsInStats", "w:doNotAutoCompressPictures",
                "w:forceUpgrade", "w:captions", "w:readModeInkLockDown",
                "w:smartTagType", "sl:schemaLibrary", "w:shapeDefaults",
                "w:doNotEmbedSmartTags", "w:decimalSymbol", "w:listSeparator",
            )
        }
        insert_at = next(
            (index for index, child in enumerate(settings) if child.tag in successors),
            len(settings),
        )
        settings.insert(insert_at, existing)
    existing.set(qn("w:val"), "true")
