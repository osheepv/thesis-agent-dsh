"""Word 原生书签、REF 域与生成后验证测试。"""

from __future__ import annotations

import zipfile

import pytest
from docx import Document

from thesis_docx.cross_reference import (
    CrossReferenceError,
    apply_cross_references,
    audit_cross_references,
    normalize_bookmark_name,
)
from thesis_docx.config import DocxConfig
from thesis_docx.generator.docx_generator import DocxGenerator
from thesis_docx.validator.docx_validator import DocxValidator
from writing.generator import SectionDraftGenerator


def test_markers_become_native_bookmark_and_ref_fields(tmp_path):
    path = tmp_path / "cross-ref.docx"
    document = Document()
    document.add_paragraph("[[BOOKMARK:TABLE-4-1|表4-1 实验结果]]")
    document.add_paragraph("结果如[[REF:TABLE-4-1|表4-1]]所示。")
    document.save(path)

    report = apply_cross_references(path)
    assert report.passed is True
    assert report.bookmark_count == 1
    assert report.reference_count == 1
    assert report.bookmark_map["TABLE-4-1"] == normalize_bookmark_name("TABLE-4-1")

    with zipfile.ZipFile(path) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        settings_xml = package.read("word/settings.xml").decode("utf-8")
    assert "bookmarkStart" in document_xml
    assert f"REF {normalize_bookmark_name('TABLE-4-1')}" in document_xml
    assert "[[BOOKMARK:" not in document_xml
    assert "[[REF:" not in document_xml
    assert "updateFields" in settings_xml
    assert audit_cross_references(path).passed is True


def test_missing_or_duplicate_cross_reference_targets_are_rejected(tmp_path):
    missing = tmp_path / "missing.docx"
    document = Document()
    document.add_paragraph("见[[REF:TABLE-9-9|表9-9]]。")
    document.save(missing)
    with pytest.raises(CrossReferenceError, match="未定义目标"):
        apply_cross_references(missing)

    duplicate = tmp_path / "duplicate.docx"
    document = Document()
    document.add_paragraph("[[BOOKMARK:TABLE-1-1|表1-1]]")
    document.add_paragraph("[[BOOKMARK:TABLE-1-1|重复表1-1]]")
    document.save(duplicate)
    with pytest.raises(CrossReferenceError, match="重复定义"):
        apply_cross_references(duplicate)


def test_docx_generator_applies_and_validator_audits_cross_references(tmp_path):
    template_path = DocxConfig.BUILTIN_TEMPLATE_PATH
    output_path = tmp_path / "output.docx"

    outcome = DocxGenerator().render(
        template_path=str(template_path),
        content={
            "chapter": (
                "[[BOOKMARK:FIGURE-3-2|图3-2 系统架构]]\n"
                "系统结构见[[REF:FIGURE-3-2|图3-2]]。"
            ),
            "topic": "原生交叉引用测试",
            "outline": "1 测试",
        },
        output_path=str(output_path),
    )
    assert outcome.cross_reference_report["passed"] is True
    assert outcome.cross_reference_report["bookmark_count"] == 1
    assert outcome.cross_reference_report["reference_count"] == 1

    validation = DocxValidator().validate(str(output_path), strict=True)
    assert validation.is_valid is True, validation.errors
    assert validation.schema_valid is True, validation.errors
    assert validation.cross_reference_valid is True
    assert validation.cross_reference_report["reference_count"] == 1
    assert validation.load_valid is True
    assert validation.roundtrip_valid is True


def test_builtin_generator_materializes_markdown_as_word_structure(tmp_path):
    output_path = tmp_path / "structured.docx"
    body = (
        "# 第1章 绪论\n\n"
        "## 1.1 研究背景\n\n"
        "第一段正文。\n\n第二段正文。\n\n"
        "# 参考文献\n\n"
        "[1] A. Example[J]. 2026.\n"
        "[2] B. Example[J]. 2025."
    )

    DocxGenerator().render(
        template_path=str(DocxConfig.BUILTIN_TEMPLATE_PATH),
        content={"chapter": body, "topic": "结构化测试", "outline": "1 绪论"},
        output_path=str(output_path),
    )

    document = Document(output_path)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    styles = {paragraph.text: paragraph.style.name for paragraph in paragraphs}
    assert styles["第1章 绪论"] == "Heading 1"
    assert styles["1.1 研究背景"] == "Heading 2"
    assert styles["参考文献"] == "Heading 1"
    assert sum(1 for paragraph in paragraphs if paragraph.text.startswith("[")) == 2


def test_section_fallback_emits_result_target_and_auditable_marker():
    generated = SectionDraftGenerator._fallback(
        {
            "title": "实验结果",
            "claims": [],
            "results": [
                {
                    "result_id": "RES-ABC123",
                    "metric": "accuracy",
                    "value": "0.93",
                    "unit": "ratio",
                    "table_or_figure_id": "TABLE-4-1",
                }
            ],
        }
    )
    assert "[[BOOKMARK:TABLE-4-1|表4-1 accuracy结果]]" in generated.content
    assert "[RES-ABC123]" in generated.content
    assert generated.used_result_ids == ["RES-ABC123"]
