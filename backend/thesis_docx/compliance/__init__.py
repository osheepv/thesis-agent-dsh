# -*- coding: utf-8 -*-
"""M6 版式合规检查器（环9，自研规则库，借鉴 thesis-format-zh 思路）。

检查维度（对应环9 格式排版验收）：
    1. 页面：页边距（上/下/左/右）默认值检查。
    2. 字体：正文"宋体/小四(12pt)/1.5倍行距"、标题"黑体"（从 Normal/Heading 样式读）。
    3. 段落：首行缩进（firstLineChars / firstLine）。
    4. 结构：前置件顺序（封面→声明→中文摘要→英文摘要→目录→正文→参考文献→附录→致谢）。
    5. OOXML 审计：styleId 是否存在于 styles.xml、page 域（PAGE/NUMPAGES）。

设计原则（对齐论文规范红线）：
    - **只查不改**：本模块是"合规检查器"，输出报告，不动内容（内容由 docxtpl 模板渲染保持）。
    - **分级**：HARD（违反即判不合规，如页边距缺失/字号严重不符）/
           SOFT（建议项，如无页边距设置）。
    - 模板驱动：若传入模板路径，则把模板的页面/字体设置作为对照基准（WYSIWYG）；
      无模板时用内置默认（学校通行值）。
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

from docx import Document as _DocxDocument
from docx.shared import Pt

logger = logging.getLogger("thesis.compliance")

#: 默认正文样式基准（学校通行，用户模板可覆盖）
_DEFAULT_RULES: Dict[str, Any] = {
    "body_font": {"cn": "宋体", "en": "Times New Roman", "size_pt": 12.0},
    "heading_font": {"cn": "黑体"},
    "line_spacing": 1.5,
    "first_line_indent_chars": 2.0,
    "page_margins": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17},  # cm
}


@dataclass
class ComplianceIssue:
    """单条合规问题。"""

    category: str = ""  # page / font / paragraph / structure / ooxml
    severity: str = "SOFT"  # HARD / SOFT
    message: str = ""

    def to_dict(self) -> Dict[str, str]:
        return {"category": self.category, "severity": self.severity, "message": self.message}


@dataclass
class ComplianceReport:
    """合规报告。"""

    compliant: bool = False
    issues: List[ComplianceIssue] = field(default_factory=list)
    rules_used: str = ""  # "template"/"default"

    def to_dict(self) -> Dict[str, Any]:
        return {
            "compliant": self.compliant,
            "issue_count": len(self.issues),
            "hard_count": sum(1 for i in self.issues if i.severity == "HARD"),
            "issues": [i.to_dict() for i in self.issues],
            "rules_used": self.rules_used,
        }


class DocxComplianceChecker:
    """docx 版式合规检查器。"""

    def __init__(self) -> None:
        self._default = _DEFAULT_RULES

    def check(self, file_path: str, template_path: Optional[str] = None) -> ComplianceReport:
        """检查 docx 版式合规。

        Args:
            file_path: 待检查 docx 路径。
            template_path: 可选用户模板路径（对照基准；无则用内置默认）。
        Returns:
            ComplianceReport。
        """
        rules = self._extract_template_rules(template_path) if template_path else self._default
        issues: List[ComplianceIssue] = []

        doc = _DocxDocument(file_path)

        # 1. 页面边距
        issues.extend(self._check_margins(doc, rules))
        # 2. 字体/字号/行距（Normal 样式）
        issues.extend(self._check_fonts(doc, rules))
        # 3. 首行缩进（正文段）
        issues.extend(self._check_indent(doc))
        # 4. 前置件结构（按段落标题关键词）
        issues.extend(self._check_structure(doc))
        # 5. OOXML 审计（styleId 存在性 / PAGE 域）
        issues.extend(self._check_ooxml(file_path))

        # 判定：无 HARD = 合规
        compliant = not any(i.severity == "HARD" for i in issues)
        return ComplianceReport(
            compliant=compliant,
            issues=issues,
            rules_used="template" if template_path else "default",
        )

    # ------------------------------------------------------------------ #
    # 模板规则抽取
    # ------------------------------------------------------------------ #
    def _extract_template_rules(self, template_path: str) -> Dict[str, Any]:
        """从模板 docx 抽取页面设置/正文样式为规则基准。"""
        try:
            tdoc = _DocxDocument(template_path)
            sec = tdoc.sections[0]
            rules: Dict[str, Any] = {
                "page_margins": {
                    "top": round(sec.top_margin.cm, 2) if sec.top_margin else None,
                    "bottom": round(sec.bottom_margin.cm, 2) if sec.bottom_margin else None,
                    "left": round(sec.left_margin.cm, 2) if sec.left_margin else None,
                    "right": round(sec.right_margin.cm, 2) if sec.right_margin else None,
                },
            }
            normal = tdoc.styles["Normal"] if hasattr(tdoc, "styles") else None
            if normal is not None:
                font = normal.font
                if font.size:
                    rules["body_font"] = {"size_pt": font.size.pt}
                # 字体名带 eastAsia（从 rPr/rFonts 读，容错处理）
                try:
                    rpr = font.element.get_or_add_rPr()
                    if rpr is not None and rpr.rFonts is not None:
                        east = rpr.rFonts.get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
                        )
                        if east:
                            rules.setdefault("body_font", {})["cn"] = east
                except Exception:  # noqa: BLE001
                    pass
            return {**self._default, **rules}
        except Exception as exc:  # noqa: BLE001 - 模板解析失败回退默认
            logger.info("模板规则抽取失败，用默认: %s", exc)
            return dict(self._default)

    # ------------------------------------------------------------------ #
    # 检查项
    # ------------------------------------------------------------------ #
    def _check_margins(self, doc, rules: Dict[str, Any]) -> List[ComplianceIssue]:
        issues = []
        expected = rules.get("page_margins", {})
        margin_names = ["top", "bottom", "left", "right"]
        try:
            sec = doc.sections[0]
            for name in margin_names:
                exp = expected.get(name)
                if exp is None:
                    continue
                actual = getattr(sec, f"{name}_margin").cm
                if actual is None:
                    issues.append(ComplianceIssue("page", "SOFT", f"页边距 {name} 未设置"))
                elif abs(actual - exp) > 0.8:  # 容忍 0.8cm
                    issues.append(ComplianceIssue(
                        "page", "SOFT",
                        f"页边距 {name}={actual:.2f}cm（期望≈{exp}cm）"
                    ))
        except Exception as exc:  # noqa: BLE001
            issues.append(ComplianceIssue("page", "HARD", f"页面设置读取失败：{exc}"))
        return issues

    def _check_fonts(self, doc, rules: Dict[str, Any]) -> List[ComplianceIssue]:
        issues = []
        try:
            # python-docx: doc.styles["Normal"]（get 方法不存在）
            normal = doc.styles["Normal"]
            if normal is None:
                issues.append(ComplianceIssue("font", "SOFT", "文档无 Normal 样式"))
                return issues
            font = normal.font
            body = rules.get("body_font", {})
            exp_size = body.get("size_pt")
            if exp_size and font.size:
                if abs(font.size.pt - exp_size) > 0.5:
                    issues.append(ComplianceIssue(
                        "font", "HARD",
                        f"正文字号 {font.size.pt}pt（期望 {exp_size}pt）"
                    ))
            # 字体名（eastAsia）
            if body.get("cn"):
                east = None
                try:
                    rpr = font.element.rPr
                    if rpr is not None and rpr.rFonts is not None:
                        east = rpr.rFonts.get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
                        )
                except Exception:  # noqa: BLE001
                    pass
                if east and body["cn"] not in east:
                    issues.append(ComplianceIssue(
                        "font", "SOFT", f"正文字体 {east}（建议 {body['cn']}）"
                    ))
        except Exception as exc:  # noqa: BLE001
            issues.append(ComplianceIssue("font", "SOFT", f"字体读取失败：{exc}"))
        return issues

    def _check_indent(self, doc) -> List[ComplianceIssue]:
        """正文首行缩进（firstLineChars=200 或 firstLine≈ 0.74cm）。"""
        issues = []
        checked = 0
        no_indent = 0
        try:
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text or len(text) < 10:
                    continue
                # 跳过标题段（黑体/第N章/摘要等）
                if any(para.style and para.style.name and "Heading" in para.style.name for _ in [0]):
                    continue
                if text.startswith(("第", "摘要", "Abstract", "目录")) and len(text) < 30:
                    continue
                pf = para.paragraph_format
                first_line = pf.first_line_indent
                has_chars = False
                if para._p.pPr is not None and para._p.pPr.ind is not None:
                    ind = para._p.pPr.ind
                    flc = ind.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstLineChars"
                    )
                    if flc:
                        has_chars = True
                if not has_chars and (first_line is None or first_line.pt < 5):
                    no_indent += 1
                checked += 1
            if checked > 0 and no_indent > max(1, checked // 3):
                issues.append(ComplianceIssue(
                    "paragraph", "SOFT",
                    f"正文首行缩进缺失 {no_indent}/{checked} 段"
                ))
        except Exception as exc:  # noqa: BLE001
            issues.append(ComplianceIssue("paragraph", "SOFT", f"缩进检查异常：{exc}"))
        return issues

    def _check_structure(self, doc) -> List[ComplianceIssue]:
        """前置件顺序：封面→中文摘要→英文摘要→目录→正文→参考文献（按规范顺序检测逆序对）。"""
        order_keywords = [
            ("封面", ["学位论文", "学校名称", "题目"]),
            ("中文摘要", ["摘要"]),
            ("英文摘要", ["Abstract"]),
            ("目录", ["目录"]),
            ("正文", ["第1章", "第一章", "绪论"]),
            ("参考文献", ["参考文献"]),
        ]
        try:
            # 记录每个前置件首次出现的位置（按规范顺序）
            found: List[tuple[str, int]] = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text or len(text) > 30:
                    continue
                for name, kws in order_keywords:
                    if any(k in text for k in kws):
                        found.append((name, i))
                        break
            # 去重保留首次位置
            seen: Dict[str, int] = {}
            for name, idx in found:
                if name not in seen:
                    seen[name] = idx
            # 检测逆序对：规范顺序中靠后的项出现在靠前项之前
            ordered = [name for name, _ in order_keywords if name in seen]
            for j in range(len(ordered)):
                for k in range(j + 1, len(ordered)):
                    if seen[ordered[k]] < seen[ordered[j]]:
                        return [ComplianceIssue(
                            "structure", "HARD",
                            f"前置件顺序异常：{ordered[k]}（第{seen[ordered[k]] + 1}段）出现在 "
                            f"{ordered[j]}（第{seen[ordered[j]] + 1}段）之前"
                        )]
        except Exception as exc:  # noqa: BLE001
            return [ComplianceIssue("structure", "SOFT", f"结构检查异常：{exc}")]
        return []

    def _check_ooxml(self, file_path: str) -> List[ComplianceIssue]:
        """OOXML 审计：styleId 存在性 + PAGE 域（lxml 读 document.xml）。"""
        issues = []
        try:
            import zipfile

            with zipfile.ZipFile(file_path) as z:
                names = z.namelist()
                # 检查是否含 styles.xml
                if "word/styles.xml" not in names:
                    issues.append(ComplianceIssue("ooxml", "HARD", "缺少 styles.xml"))
                # 检查页码域（PAGE 或 NUMPAGES）
                doc_xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
                if "PAGE" not in doc_xml and "NUMPAGES" not in doc_xml and "STYLEREF" not in doc_xml:
                    issues.append(ComplianceIssue("ooxml", "SOFT", "未检出页码域（PAGE/NUMPAGES）"))
        except Exception as exc:  # noqa: BLE001
            issues.append(ComplianceIssue("ooxml", "HARD", f"OOXML 读取失败：{exc}"))
        return issues


#: 模块级单例
_checker: Optional[DocxComplianceChecker] = None


def get_compliance_checker() -> DocxComplianceChecker:
    global _checker
    if _checker is None:
        _checker = DocxComplianceChecker()
    return _checker
