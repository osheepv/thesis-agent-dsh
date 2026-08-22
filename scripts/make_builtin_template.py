# -*- coding: utf-8 -*-
"""生成内置论文模板（含标准论文占位符）。

用途：
    聚合层 docx/generate 在「无用户模板」时，需要兜底到一个带标准论文占位符的
    模板，而不是 python-docx 内置的空 default.docx。本脚本用 python-docx 在内存
    生成一个含 `{{ title }}`、`{{ degree }}`、`{{ subject_field }}`、`{{ outline }}`、
    `{{ chapter }}` 占位符的模板，并落盘到 DocxConfig.BUILTIN_TEMPLATE_PATH。

再生成：
    backend/.venv_test/Scripts/python.exe scripts/make_builtin_template.py
"""
from __future__ import annotations

import sys
from pathlib import Path

# 允许从项目根直接运行（保证 backend 包可导入）。
ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402
from docx.shared import Pt  # noqa: E402

from backend.thesis_docx.config import DocxConfig  # noqa: E402


def build() -> Document:
    """构造内置论文模板文档。"""
    doc = Document()

    # 标题
    title = doc.add_heading(level=0)
    title.add_run("{{ title }}")

    # 基本信息（学位、学科方向占位符）
    degree = doc.add_paragraph()
    degree.add_run("学位类型：").bold = True
    degree.add_run("{{ degree }}")
    subject = doc.add_paragraph()
    subject.add_run("学科方向：").bold = True
    subject.add_run("{{ subject_field }}")

    # 摘要段（可选占位符，若 content 未提供则在渲染时留空）
    doc.add_heading("摘要", level=1)
    abstract = doc.add_paragraph("{{ abstract }}")

    # 大纲段
    doc.add_heading("论文大纲", level=1)
    outline = doc.add_paragraph()
    outline.add_run("{{ outline }}")

    # 正文段（核心章节内容）
    doc.add_heading("正文", level=1)
    chapter = doc.add_paragraph()
    chapter.add_run("{{ chapter }}")

    # 设置正文样式（宋体/Times，适度即可，docxtpl 只关心占位符）。
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    return doc


def main() -> None:
    """生成模板并落盘。"""
    target = DocxConfig.BUILTIN_TEMPLATE_PATH
    target.parent.mkdir(parents=True, exist_ok=True)
    build().save(str(target))
    print(f"built-in template saved -> {target}")


if __name__ == "__main__":
    main()
